"""
LexGuard AI — Document Processor
Handles PDF and DOCX text extraction. Uses PyMuPDF for PDFs, python-docx for Word docs.
"""

import os
from typing import Optional
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from config import UPLOAD_DIR, ALLOWED_EXTENSIONS


class DocumentProcessor:
    """Extract text from uploaded legal documents."""

    @staticmethod
    def validate_file(filename: str) -> bool:
        """Check if file extension is allowed."""
        ext = os.path.splitext(filename)[1].lower()
        return ext in ALLOWED_EXTENSIONS

    @staticmethod
    async def extract_text(file_path: str) -> str:
        """
        Extract all text from PDF or DOCX file.
        Returns plain text with basic formatting preserved (newlines, spacing).
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return DocumentProcessor._extract_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentProcessor._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(file_path)
        text_blocks = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text_blocks.append(f"[Page {page_num + 1}]\n")
            text_blocks.append(page.get_text("text"))  # plain text mode
            text_blocks.append("\n\n")

        doc.close()
        return "".join(text_blocks).strip()

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @staticmethod
    def get_document_metadata(file_path: str) -> dict:
        """Extract page count, file size, etc."""
        ext = os.path.splitext(file_path)[1].lower()
        size_bytes = os.path.getsize(file_path)

        metadata = {
            "file_size_bytes": size_bytes,
            "file_size_mb": round(size_bytes / (1024 * 1024), 2),
            "extension": ext,
        }

        if ext == ".pdf":
            doc = fitz.open(file_path)
            metadata["page_count"] = len(doc)
            doc.close()
        elif ext in [".docx", ".doc"]:
            doc = DocxDocument(file_path)
            # Approximate page count (250 words per page)
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            metadata["page_count"] = max(1, word_count // 250)

        return metadata
